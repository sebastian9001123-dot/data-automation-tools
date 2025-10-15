# manuals.py
import os
import copy
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# ----------------------------------------------------------------
# UTILIDADES: copiar contenido y forzar bordes en tablas
# ----------------------------------------------------------------

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

def ensure_table_borders(tbl_element):
    """
    Recibe un elemento <w:tbl> (XML) y asegura que tenga tblPr/tblBorders
    y que las celdas tengan tcPr/tcBorders (incluye insideH/insideV).
    """
    # tblPr
    tblPr = tbl_element.find("w:tblPr", namespaces=NS)
    if tblPr is None:
        tblPr = parse_xml(f"<w:tblPr {nsdecls('w')}/>")
        tbl_element.insert(0, tblPr)

    # Crear tblBorders (reemplaza si existía)
    tblBorders_xml = f"""
    <w:tblBorders {nsdecls('w')}>
      <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
    </w:tblBorders>
    """
    tblBorders = parse_xml(tblBorders_xml)
    old = tblPr.find("w:tblBorders", namespaces=NS)
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)

    # Aplicar bordes a cada celda (tcPr/tcBorders)
    for tc in tbl_element.findall(".//w:tc", namespaces=NS):
        tcPr = tc.find("w:tcPr", namespaces=NS)
        if tcPr is None:
            tcPr = parse_xml(f"<w:tcPr {nsdecls('w')}/>")
            tc.insert(0, tcPr)
        tcBorders_xml = f"""
        <w:tcBorders {nsdecls('w')}>
          <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
          <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
          <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
          <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tcBorders>
        """
        tcBorders = parse_xml(tcBorders_xml)
        old_tc = tcPr.find("w:tcBorders", namespaces=NS)
        if old_tc is not None:
            tcPr.remove(old_tc)
        tcPr.append(tcBorders)

def copy_doc_contents_preserve_tables(src_doc: Document, dest_doc: Document):
    """
    Copia elementos del documento src_doc al dest_doc.
    Para tablas, aplica bordes y mantiene su estructura.
    """
    for element in src_doc.element.body:
        new_el = copy.deepcopy(element)

        # Si tiene tag que contiene 'tbl' (tabla), forzamos bordes
        # Nota: new_el.tag puede ser algo como '{...}tbl'
        if "tbl" in new_el.tag:
            # new_el es la tabla completa; asegurar bordes en tabla y celdas
            ensure_table_borders(new_el)

        dest_doc.element.body.append(new_el)

def replace_text_in_doc(doc: Document, replacements: dict):
    """Reemplaza texto dentro de párrafos y tablas (reconstruyendo runs para evitar fragmentación)."""
    # Párrafos
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                new_text = p.text.replace(key, str(val))
                # eliminar runs y escribir nuevo
                for run in list(p.runs):
                    p._element.remove(run._element)
                p.add_run(new_text)

    # Tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        # reemplazo directo en celda
                        cell.text = cell.text.replace(key, str(val))

# ----------------------------------------------------------------
# INTERFAZ GRÁFICA
# ----------------------------------------------------------------

class GeneradorGUI:
    def __init__(self, root):
        self.root = root
        root.title("MANUALS versión 1.0")
        root.geometry("850x640")

        self.membretes = {}   # asociación -> ruta
        self.manuales = []    # rutas
        self.trabajadores = []
        self.carpeta_salida = None

        # Título
        title_frame = ttk.Frame(root, padding=(8, 5))
        title_frame.pack(fill="x")
        ttk.Label(title_frame, text="🗂️  MANUALS versión 1.0", font=("Segoe UI", 16, "bold"), foreground="#0052cc").pack(side="left", padx=5)
        ttk.Label(title_frame, text="(Gestor de membretes y manuales automáticos)", font=("Segoe UI", 10), foreground="#555").pack(side="left", padx=10)

        # Membretes
        frm_top = ttk.Frame(root, padding=8)
        frm_top.pack(fill="x")

        ttk.Label(frm_top, text="Membretes (asigna uno por asociación):").grid(row=0, column=0, sticky="w")
        frm_memb_btns = ttk.Frame(frm_top)
        frm_memb_btns.grid(row=0, column=1, sticky="e")
        ttk.Button(frm_memb_btns, text="Agregar membrete", command=self.agregar_membrete).pack(side="left", padx=3)
        ttk.Button(frm_memb_btns, text="Eliminar seleccionado", command=self.eliminar_membrete).pack(side="left", padx=3)
        self.lst_membretes = tk.Listbox(frm_top, height=5)
        self.lst_membretes.grid(row=1, column=0, columnspan=2, sticky="we", pady=5)

        # Manuales
        ttk.Label(frm_top, text="Manuales base (.docx):").grid(row=2, column=0, sticky="w", pady=(8, 0))
        frm_manual_btns = ttk.Frame(frm_top)
        frm_manual_btns.grid(row=2, column=1, sticky="e")
        ttk.Button(frm_manual_btns, text="Agregar manual(es)", command=self.agregar_manuales).pack(side="left", padx=3)
        ttk.Button(frm_manual_btns, text="Eliminar seleccionado", command=self.eliminar_manual).pack(side="left", padx=3)
        self.lst_manuales = tk.Listbox(frm_top, height=4)
        self.lst_manuales.grid(row=3, column=0, columnspan=2, sticky="we", pady=5)

        # Documentos a generar
        frm_mid = ttk.Frame(root, padding=8)
        frm_mid.pack(fill="both", expand=True)
        ttk.Label(frm_mid, text="Documentos a generar (Nombre del documento, Cargo, Asociación, Manual):").pack(anchor="w")
        cols = ("Nombre documento", "Cargo", "Asociacion", "Manual")
        self.tree = ttk.Treeview(frm_mid, columns=cols, show="headings", height=6)
        for c in cols:
            self.tree.heading(c, text=c)
            self.tree.column(c, width=180)
        self.tree.pack(fill="both", expand=True, pady=5)
        frm_buttons = ttk.Frame(frm_mid)
        frm_buttons.pack(fill="x")
        ttk.Button(frm_buttons, text="Agregar documento", command=self.agregar_trabajador).pack(side="left", padx=4)
        ttk.Button(frm_buttons, text="Editar seleccionado", command=self.editar_trabajador).pack(side="left", padx=4)
        ttk.Button(frm_buttons, text="Eliminar seleccionado", command=self.eliminar_trabajador).pack(side="left", padx=4)
        ttk.Button(frm_buttons, text="Limpiar lista", command=self.limpiar_trabajadores).pack(side="left", padx=4)

        # Carpeta salida y generar
        frm_bot = ttk.Frame(root, padding=8)
        frm_bot.pack(fill="x")
        ttk.Button(frm_bot, text="Seleccionar carpeta de salida", command=self.seleccionar_salida).pack(side="left")
        self.salidadir_label = ttk.Label(frm_bot, text="(ninguna carpeta seleccionada)")
        self.salidadir_label.pack(side="left", padx=8)
        ttk.Button(frm_bot, text="Generar documentos", command=self.generar_documentos).pack(side="right")

        # Footer creador
        footer = ttk.Frame(root, padding=5)
        footer.pack(fill="x", side="bottom")
        ttk.Label(footer, text="Creado por: Sebastián Ramírez", font=("Segoe UI", 9), foreground="#777").pack(side="left", padx=5)

    # --- funciones de manejo ---
    def agregar_membrete(self):
        rutas = filedialog.askopenfilenames(title="Selecciona membrete(s) .docx", filetypes=[("Word", "*.docx")])
        for r in rutas:
            if not r:
                continue
            nombre_asoc = simpledialog.askstring("Asociación", f"Nombre o clave de la asociación para este membrete:\n\nArchivo: {os.path.basename(r)}")
            if not nombre_asoc:
                continue
            self.membretes[nombre_asoc] = r
            self.lst_membretes.insert(tk.END, f"{nombre_asoc} - {os.path.basename(r)}")

    def eliminar_membrete(self):
        sel = self.lst_membretes.curselection()
        if not sel:
            messagebox.showinfo("Eliminar membrete", "Selecciona un membrete de la lista.")
            return
        index = sel[0]
        texto = self.lst_membretes.get(index)
        clave = texto.split(" - ")[0].strip()
        if clave in self.membretes:
            del self.membretes[clave]
        self.lst_membretes.delete(index)

    def agregar_manuales(self):
        rutas = filedialog.askopenfilenames(title="Selecciona manual(es) base", filetypes=[("Word", "*.docx")])
        for r in rutas:
            if r and r not in self.manuales:
                self.manuales.append(r)
                self.lst_manuales.insert(tk.END, os.path.basename(r))

    def eliminar_manual(self):
        sel = self.lst_manuales.curselection()
        if not sel:
            messagebox.showinfo("Eliminar manual", "Selecciona un manual de la lista.")
            return
        index = sel[0]
        del self.manuales[index]
        self.lst_manuales.delete(index)

    def agregar_trabajador(self):
        if not self.manuales:
            messagebox.showerror("Falta manual", "Primero debes cargar al menos un manual base.")
            return
        win = tk.Toplevel(self.root)
        win.title("Agregar documento")
        ttk.Label(win, text="Nombre del documento generado:").grid(row=0, column=0, sticky="e")
        ent_nombre = ttk.Entry(win, width=40); ent_nombre.grid(row=0, column=1, pady=4)
        ttk.Label(win, text="Cargo:").grid(row=1, column=0, sticky="e")
        ent_cargo = ttk.Entry(win, width=40); ent_cargo.grid(row=1, column=1, pady=4)
        ttk.Label(win, text="Asociación (clave):").grid(row=2, column=0, sticky="e")
        ent_asoc = ttk.Entry(win, width=40); ent_asoc.grid(row=2, column=1, pady=4)
        ttk.Label(win, text="Selecciona manual:").grid(row=3, column=0, sticky="e")
        cmb_manual = ttk.Combobox(win, values=[os.path.basename(x) for x in self.manuales], state="readonly", width=37)
        cmb_manual.current(0); cmb_manual.grid(row=3, column=1, pady=4)
        def on_ok():
            nombre_doc = ent_nombre.get().strip(); cargo = ent_cargo.get().strip(); asoc = ent_asoc.get().strip(); manual_idx = cmb_manual.current()
            if not nombre_doc or not asoc:
                messagebox.showerror("Datos incompletos", "El nombre del documento y la asociación son obligatorios.")
                return
            item = {'nombre_doc': nombre_doc, 'cargo': cargo, 'asociacion': asoc, 'manual_idx': manual_idx}
            self.trabajadores.append(item)
            self.tree.insert("", tk.END, values=(nombre_doc, cargo, asoc, os.path.basename(self.manuales[manual_idx])))
            win.destroy()
        ttk.Button(win, text="Agregar", command=on_ok).grid(row=4, column=1, sticky="e", pady=8)

    def editar_trabajador(self):
        sel = self.tree.selection()
        if not sel:
            messagebox.showinfo("Selecciona", "Selecciona un documento para editar.")
            return
        idx = self.tree.index(sel[0]); reg = self.trabajadores[idx]
        win = tk.Toplevel(self.root); win.title("Editar documento")
        ttk.Label(win, text="Nombre del documento generado:").grid(row=0, column=0, sticky="e")
        ent_nombre = ttk.Entry(win, width=40); ent_nombre.insert(0, reg['nombre_doc']); ent_nombre.grid(row=0, column=1, pady=4)
        ttk.Label(win, text="Cargo:").grid(row=1, column=0, sticky="e"); ent_cargo = ttk.Entry(win, width=40); ent_cargo.insert(0, reg['cargo']); ent_cargo.grid(row=1, column=1, pady=4)
        ttk.Label(win, text="Asociación (clave):").grid(row=2, column=0, sticky="e"); ent_asoc = ttk.Entry(win, width=40); ent_asoc.insert(0, reg['asociacion']); ent_asoc.grid(row=2, column=1, pady=4)
        ttk.Label(win, text="Selecciona manual:").grid(row=3, column=0, sticky="e")
        cmb_manual = ttk.Combobox(win, values=[os.path.basename(x) for x in self.manuales], state="readonly", width=37)
        cmb_manual.current(reg['manual_idx']); cmb_manual.grid(row=3, column=1, pady=4)
        def on_ok():
            nombre_doc = ent_nombre.get().strip(); cargo = ent_cargo.get().strip(); asoc = ent_asoc.get().strip(); manual_idx = cmb_manual.current()
            if not nombre_doc or not asoc:
                messagebox.showerror("Datos incompletos", "El nombre del documento y la asociación son obligatorios.")
                return
            self.trabajadores[idx] = {'nombre_doc': nombre_doc, 'cargo': cargo, 'asociacion': asoc, 'manual_idx': manual_idx}
            self.tree.item(sel[0], values=(nombre_doc, cargo, asoc, os.path.basename(self.manuales[manual_idx])))
            win.destroy()
        ttk.Button(win, text="Guardar", command=on_ok).grid(row=4, column=1, sticky="e", pady=8)

    def eliminar_trabajador(self):
        sel = self.tree.selection()
        if not sel:
            return
        idx = self.tree.index(sel[0]); self.tree.delete(sel[0]); del self.trabajadores[idx]

    def limpiar_trabajadores(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        self.trabajadores = []

    def seleccionar_salida(self):
        carpeta = filedialog.askdirectory(title="Selecciona carpeta de salida")
        if carpeta:
            self.carpeta_salida = carpeta
            self.salidadir_label.config(text=self.carpeta_salida)

    def generar_documentos(self):
        if not self.trabajadores:
            messagebox.showerror("Nada que generar", "No hay documentos en la lista.")
            return
        if not self.carpeta_salida:
            messagebox.showerror("Carpeta salida", "Selecciona una carpeta de salida.")
            return
        errores = []; contador = 0
        for reg in self.trabajadores:
            nombre_doc = reg['nombre_doc']; cargo = reg['cargo']; asoc = reg['asociacion']; manual_idx = reg['manual_idx']
            if asoc not in self.membretes:
                errores.append(f"No hay membrete para la asociación '{asoc}'"); continue
            ruta_membrete = self.membretes[asoc]; ruta_manual = self.manuales[manual_idx]
            try:
                doc_membrete = Document(ruta_membrete)
                doc_manual = Document(ruta_manual)
                # copiar contenido y asegurarse bordes en tablas
                copy_doc_contents_preserve_tables(doc_manual, doc_membrete)
                # reemplazos
                replace_text_in_doc(doc_membrete, {"{{CARGO}}": cargo, "{{ASOCIACION}}": asoc})
                salida = os.path.join(self.carpeta_salida, f"{nombre_doc}.docx")
                doc_membrete.save(salida)
                contador += 1
            except Exception as e:
                errores.append(f"Error con '{nombre_doc}': {e}")
        msg = f"Generación completada. Documentos creados: {contador}/{len(self.trabajadores)}."
        if errores:
            msg += "\n\nProblemas:\n" + "\n".join(errores[:10])
            messagebox.showwarning("Completado con errores", msg)
        else:
            messagebox.showinfo("Éxito", msg)

# ----------------------------------------------------------------
# EJECUCIÓN
# ----------------------------------------------------------------

if __name__ == "__main__":
    root = tk.Tk()
    app = GeneradorGUI(root)
    root.mainloop()
